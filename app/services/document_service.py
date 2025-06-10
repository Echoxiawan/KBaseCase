import os
import io
import tempfile
# import pytesseract
from PIL import Image
from PyPDF2 import PdfReader
from docx import Document
import markdown
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pdf2image import convert_from_path
import zipfile
import xml.etree.ElementTree as ET
from dotenv import load_dotenv
from paddleocr import PaddleOCR
import time
from ..utils.logger import get_logger

# 获取日志器
logger = get_logger('document_service')

# 加载环境变量
load_dotenv()


class DocumentService:
    """文档处理服务，负责解析不同格式的文档并提取文本内容，包括图片中的文本"""
    
    # 单例实例
    _instance = None
    
    def __new__(cls):
        """实现单例模式，确保只创建一个DocumentService实例"""
        if cls._instance is None:
            cls._instance = super(DocumentService, cls).__new__(cls)
            cls._instance._initialized = False
        return cls._instance

    def __init__(self):
        """初始化方法，只在第一次创建实例时执行"""
        # 如果已经初始化过，则直接返回
        if getattr(self, '_initialized', False):
            return
            
        # 检查OCR总开关是否启用
        ocr_enabled = os.environ.get('OCR_ENABLED', 'True').lower() == 'true'
        
        # 如果OCR被禁用，直接设置tesseract_available为False并跳过初始化
        if not ocr_enabled:
            logger.info("OCR功能已禁用，将跳过所有OCR处理")
            # 仅在真正初始化时打印信息，而不是每次创建实例
            self.tesseract_available = False
            self.ocr_engine = None
            self._initialized = True
            return
            
        # 初始化PaddleOCR，替代原来的Tesseract
        try:
            logger.info("正在初始化PaddleOCR...")
            start_time = time.time()

            # 从环境变量中获取OCR配置参数
            lang = os.environ.get('OCR_LANG', 'ch')
            # 初始化PaddleOCR实例
            self.ocr_engine = PaddleOCR(use_angle_cls=True, lang=lang)
            # 为了兼容性，保留tesseract_available标志，但始终设为True
            self.tesseract_available = True

            elapsed_time = time.time() - start_time
            logger.info(f"PaddleOCR初始化成功，耗时: {elapsed_time:.2f}秒")
            # 仅在真正初始化时打印信息，而不是每次创建实例
        except Exception as e:
            logger.error(f"PaddleOCR初始化失败: {str(e)}", exc_info=True)
            # 仅在真正初始化时打印信息，而不是每次创建实例
            self.tesseract_available = False
            self.ocr_engine = None
        
        # 标记为已初始化
        self._initialized = True

    def process_document(self, file_path, file_type):
        """处理文档并返回文本内容"""
        if file_type == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            return self.extract_text_from_docx(file_path)
        elif file_type == 'md':
            return self.extract_text_from_markdown(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")

    def extract_text_from_pdf(self, file_path):
        """从PDF文件中提取文本，包括图片中的文本"""
        text = ""

        # 使用PyPDF2提取文本
        with open(file_path, 'rb') as f:
            pdf_reader = PdfReader(f)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() or ""

        # 如果 OCR引擎 不可用，跳过 OCR 处理
        if not self.tesseract_available:
            return text

        # 使用pdf2image提取图片并进行OCR
        try:
            # 创建临时目录存储所有图片
            with tempfile.TemporaryDirectory() as temp_dir:
                # 将PDF页面转换为图片
                images = convert_from_path(file_path)

                # 对每个图片进行OCR
                for i, image in enumerate(images):
                    # 在临时目录中创建图片文件
                    img_path = os.path.join(temp_dir, f"page_{i + 1}.png")
                    image.save(img_path)

                    # 提取图片中的文本
                    img_text = self.extract_text_from_image(img_path)
                    if img_text:
                        text += f"\n[图片文本 - 页面 {i + 1}]\n{img_text}\n"

                    # 无需手动删除文件，临时目录会自动清理
        except Exception as e:
            logger.error(f"PDF图片处理失败: {str(e)}", exc_info=True)

        return text

    def extract_text_from_docx(self, file_path):
        """从DOCX文件中提取文本，包括图片中的文本"""
        text = ""

        # 提取文本内容
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"

        # 从表格中提取文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        # 如果 OCR引擎 不可用，跳过 OCR 处理
        if not self.tesseract_available:
            return text

        # 提取DOCX中的图片并进行OCR
        try:
            # 创建临时目录存储所有图片
            with tempfile.TemporaryDirectory() as temp_dir:
                # DOCX文件实际上是一个ZIP文件
                with zipfile.ZipFile(file_path) as docx_zip:
                    # 获取文档关系文件
                    rels_content = docx_zip.read('word/_rels/document.xml.rels')
                    rels_root = ET.fromstring(rels_content)

                    # 查找所有图片关系
                    image_rels = {}
                    for rel in rels_root.findall(
                            './/{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                        if 'image' in rel.get('Target', ''):
                            image_rels[rel.get('Id')] = rel.get('Target')

                    # 提取图片并进行OCR
                    for rel_id, target in image_rels.items():
                        # 修正图片路径
                        image_path = f"word/{target.replace('../', '')}"

                        try:
                            # 读取图片数据
                            image_data = docx_zip.read(image_path)

                            # 在临时目录中创建图片文件
                            img_file_name = os.path.basename(image_path)
                            img_temp_path = os.path.join(temp_dir, img_file_name)

                            # 写入图片数据到临时文件
                            with open(img_temp_path, 'wb') as img_file:
                                img_file.write(image_data)

                            # 提取图片中的文本
                            img_text = self.extract_text_from_image(img_temp_path)
                            if img_text:
                                text += f"\n[图片文本 - {img_file_name}]\n{img_text}\n"

                            # 无需手动删除文件，临时目录会自动清理
                        except Exception as e:
                            logger.error(f"DOCX图片处理失败 ({image_path}): {str(e)}", exc_info=True)
        except Exception as e:
            logger.error(f"DOCX图片提取失败: {str(e)}", exc_info=True)

        return text

    def extract_text_from_markdown(self, file_path):
        """从Markdown文件中提取文本"""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # 将Markdown转换为HTML，然后去除HTML标签
        html = markdown.markdown(md_content)
        # 简单地去除HTML标签（更复杂的场景可能需要使用BeautifulSoup）
        text = html.replace('<p>', '').replace('</p>', '\n')
        for tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            text = text.replace(f'<{tag}>', '').replace(f'</{tag}>', '\n')
        text = text.replace('<li>', '- ').replace('</li>', '\n')
        text = text.replace('<ul>', '').replace('</ul>', '\n')
        text = text.replace('<ol>', '').replace('</ol>', '\n')

        # 如果 OCR引擎 不可用，跳过 OCR 处理
        if not self.tesseract_available:
            return text

        # 检查Markdown中是否包含图片链接
        img_links = []
        lines = md_content.split('\n')
        for line in lines:
            if '![' in line and '](' in line and ')' in line:
                # 提取图片链接
                start_idx = line.find('](') + 2
                end_idx = line.find(')', start_idx)
                if start_idx > 1 and end_idx > start_idx:
                    img_link = line[start_idx:end_idx]
                    img_links.append(img_link)

        # 如果存在本地图片链接，尝试进行OCR
        for img_link in img_links:
            if not img_link.startswith(('http://', 'https://')):
                # 假设是相对路径，尝试从文档所在目录解析
                img_path = os.path.join(os.path.dirname(file_path), img_link)
                if os.path.exists(img_path):
                    img_text = self.extract_text_from_image(img_path)
                    if img_text:
                        text += f"\n[图片文本 - {os.path.basename(img_link)}]\n{img_text}\n"

        return text

    def extract_text_from_image(self, image_path):
        """从图片中提取文本使用OCR"""
        # 如果 OCR引擎 不可用，直接返回空字符串
        if not self.tesseract_available:
            logger.warning(f"OCR引擎不可用，跳过图片处理: {image_path}")
            return ""

        # 获取最大重试次数
        max_retries = int(os.environ.get('OCR_MAX_RETRIES', '3'))
        retry_count = 0

        logger.info(f"开始处理图片: {image_path}")
        start_time = time.time()
        while retry_count < max_retries:
            try:
                # 使用PaddleOCR进行OCR识别，移除cls参数
                logger.debug(f"调用OCR引擎识别图片: {image_path}")
                result = self.ocr_engine.predict(image_path)

                # 处理OCR结果
                if not result or len(result) == 0:
                    logger.info(f"OCR未识别出文本: {image_path}")
                    return ""

                # 设置置信度阈值，可以通过环境变量配置
                confidence_threshold = float(os.environ.get('OCR_CONFIDENCE_THRESHOLD', '0.5'))
                logger.debug(f"OCR置信度阈值: {confidence_threshold}")

                # PaddleOCR结果格式为列表，包含识别出的每一行文本信息
                texts = []
                total_items = 0
                filtered_items = 0

                for line in result:
                    for box_info in line:
                        total_items += 1
                        # box_info结构为：[[[x1,y1],[x2,y2],[x3,y3],[x4,y4]], [text, confidence]]
                        if isinstance(box_info, list) and len(box_info) == 2:
                            text = box_info[1][0]  # 提取识别的文本
                            confidence = box_info[1][1]  # 提取置信度

                            # 只保留置信度高于阈值的文本
                            if confidence >= confidence_threshold:
                                texts.append(text)
                            else:
                                filtered_items += 1
                                logger.debug(f"文本因置信度低被过滤: '{text}' (置信度: {confidence:.2f})")

                # 统计处理结果
                elapsed_time = time.time() - start_time
                logger.info(f"OCR处理完成: {image_path}, 耗时: {elapsed_time:.2f}秒, "
                            f"识别项: {total_items}, 过滤项: {filtered_items}, 保留项: {len(texts)}")

                # 合并所有文本，用换行符分隔
                text = "\n".join(texts)
                return text.strip()

            except Exception as e:
                retry_count += 1
                logger.warning(f"OCR处理失败 (尝试 {retry_count}/{max_retries}): {str(e)}", exc_info=True)
                if retry_count < max_retries:
                    # 等待一段时间后重试，使用全局导入的time模块
                    time.sleep(1)
                else:
                    logger.error(f"OCR处理最终失败，已达到最大重试次数: {str(e)}", exc_info=True)
                    return ""
        return None

    def split_text_for_context_window(self, text, max_tokens=4000, chunk_size=1000, chunk_overlap=200):
        """将文本分割为适合上下文窗口大小的块"""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=lambda x: len(x.split())  # 简单估算token数量
        )

        chunks = text_splitter.split_text(text)

        # 确保每个块不超过最大token数
        result_chunks = []
        current_chunk = ""
        current_tokens = 0

        for chunk in chunks:
            chunk_tokens = len(chunk.split())
            if current_tokens + chunk_tokens <= max_tokens:
                current_chunk += chunk
                current_tokens += chunk_tokens
            else:
                if current_chunk:
                    result_chunks.append(current_chunk)
                current_chunk = chunk
                current_tokens = chunk_tokens

        if current_chunk:
            result_chunks.append(current_chunk)

        return result_chunks

    def _check_tesseract_available(self):
        """检查 OCR 引擎是否可用（兼容性方法）"""
        # 这个方法现在只是简单返回是否已初始化PaddleOCR
        return self.ocr_engine is not None
